#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将Markdown格式的实验报告转换为Word文档（基础版本，无自动编号、无目录页）
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.parser import OxmlElement
from docx.oxml.ns import qn
import re
import os

def set_run_font(run, font_name='SimSun', font_size=12):
    """设置字体"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)

def add_hyperlink(paragraph, text, url):
    """添加超链接"""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    
    paragraph._p.append(hyperlink)
    return hyperlink

def convert_markdown_to_word(md_file, docx_file):
    """将Markdown文件转换为Word文档（基础版本）"""
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # 读取Markdown文件
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # 分割内容
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # 跳过空行
        if not line:
            i += 1
            continue
        
        # 处理标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            title_text = line.lstrip('#').strip()
            
            if level == 1:
                # 一级标题
                heading = doc.add_heading(title_text, level=1)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in heading.runs:
                    set_run_font(run, 'SimHei', 18)
                    run.bold = True
            elif level == 2:
                # 二级标题
                heading = doc.add_heading(title_text, level=2)
                for run in heading.runs:
                    set_run_font(run, 'SimHei', 16)
                    run.bold = True
            elif level == 3:
                # 三级标题
                heading = doc.add_heading(title_text, level=3)
                for run in heading.runs:
                    set_run_font(run, 'SimHei', 14)
                    run.bold = True
            elif level == 4:
                # 四级标题
                heading = doc.add_heading(title_text, level=4)
                for run in heading.runs:
                    set_run_font(run, 'SimHei', 12)
                    run.bold = True
        
        # 处理表格
        elif line.startswith('|'):
            # 收集表格行
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) >= 2:
                # 解析表格
                table_data = []
                for table_line in table_lines:
                    if table_line.startswith('|') and table_line.endswith('|'):
                        # 移除首尾的|并分割
                        cells = [cell.strip() for cell in table_line[1:-1].split('|')]
                        table_data.append(cells)
                
                if table_data:
                    # 创建表格
                    rows = len(table_data)
                    cols = len(table_data[0])
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(table_data):
                        for col_idx, cell_text in enumerate(row_data):
                            if row_idx < rows and col_idx < cols:
                                cell = table.cell(row_idx, col_idx)
                                cell.text = cell_text
                                # 设置单元格字体
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        set_run_font(run, 'SimSun', 10)
            
            continue  # 因为i已经在循环中增加了
        
        # 处理代码块
        elif line.startswith('```'):
            # 收集代码块内容
            code_lines = []
            i += 1  # 跳过开始标记
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            if code_lines:
                # 添加代码块
                code_para = doc.add_paragraph()
                code_para.style = 'No Spacing'
                code_run = code_para.add_run('\n'.join(code_lines))
                set_run_font(code_run, 'Courier New', 10)
                code_para.paragraph_format.left_indent = Inches(0.5)
                code_para.paragraph_format.space_before = Pt(6)
                code_para.paragraph_format.space_after = Pt(6)
        
        # 处理列表
        elif line.startswith('- ') or line.startswith('* '):
            # 处理无序列表
            list_text = line[2:].strip()
            para = doc.add_paragraph()
            para.style = 'List Bullet'
            run = para.add_run(list_text)
            set_run_font(run, 'SimSun', 12)
        
        elif re.match(r'^\d+\.', line):
            # 处理有序列表
            list_text = re.sub(r'^\d+\.\s*', '', line)
            para = doc.add_paragraph()
            para.style = 'List Number'
            run = para.add_run(list_text)
            set_run_font(run, 'SimSun', 12)
        
        # 处理粗体和斜体
        elif '**' in line or '*' in line:
            para = doc.add_paragraph()
            # 简单的粗体处理
            text = line
            text = text.replace('**', '')  # 移除粗体标记
            run = para.add_run(text)
            set_run_font(run, 'SimSun', 12)
        
        # 处理普通段落
        else:
            para = doc.add_paragraph()
            run = para.add_run(line)
            set_run_font(run, 'SimSun', 12)
        
        i += 1
    
    # 保存文档
    doc.save(docx_file)
    print(f"Word文档已保存为: {docx_file}")

def main():
    """主函数"""
    md_file = 'GNN_拓扑还原实验报告.md'
    docx_file = 'GNN_拓扑还原实验报告.docx'
    
    if not os.path.exists(md_file):
        print(f"错误: 找不到文件 {md_file}")
        return
    
    try:
        convert_markdown_to_word(md_file, docx_file)
        print("转换完成！")
    except Exception as e:
        print(f"转换过程中出现错误: {e}")

if __name__ == "__main__":
    main() 